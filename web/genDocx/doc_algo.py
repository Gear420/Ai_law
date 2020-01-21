

#关于上传form的设计:
#sex:0 or 1
#name:
#childnum:
#childtoparent:1 0
#tanwangfangshi:0,1 #0探望 #1短期生活
#housenum:
#fangdai:
#newsaving:0 1
#othersaving 0 1
#gongtongzhaiwu 0 1


#todo:是否有汽车

def main_algo_generate(form,name):
    from docx import Document
    from docx.oxml.ns import qn


    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    if form['sex'] == 1:
        paragraph = document.add_paragraph('甲方(男方)姓名:{} 性别:男，身份证号码：住 址:'.format(form["name"]))
        paragraph = document.add_paragraph('乙方(女方)姓名: 性别女 ，身份证号码：住 址:'.format(form["name"]))
    else:
        paragraph = document.add_paragraph('甲方(女方)姓名:{} 性别:女，身份证号码：住 址:'.format(form["name"]))
        paragraph = document.add_paragraph('乙方(男方)姓名: 性别男 ，身份证号码：住 址:'.format(form["name"]))

    paragraph = document.add_paragraph('甲乙双方于       年     月     日在       市民政局登记结婚，甲乙双方无财产、无子女。现因                     ，导致夫'
                                       '妻感情完全破裂，已无法继续共同生活，故双方自愿离婚并达成以下协议：')
    paragraph = document.add_paragraph('第一条 甲乙双方自愿协议离婚。')
    paragraph = document.add_paragraph('第二条 子女抚养')

    for i in range(form['childnum']):
        paragraph = document.add_paragraph('基于不影响子女生活和学习的原则，双方本着相互体谅的态度，协助办理如下各项事宜：')
        if form['tanwangfangshi'] == 0:
            paragraph = document.add_paragraph('婚后于___（日期）出生的子女由男方抚养。女方每月支付抚养费___元；支付时间是'
                                                   '每月___日之前支付下个月的抚养费；支付方式为银行转账，女方接受抚养费的账号'
                                                   '为:___；直至付到___周岁止，___周岁之后的有关费用双方日后重新协商。')
            paragraph = document.add_paragraph('女方每___（周/月等）有在保证子女身心健康的基础上探视子女一次的权利。男方应'
                                                   '为女方的探视提供必要的协助。')

            paragraph = document.add_paragraph('女方每___（周/月等）有在保证子女身心健康的基础上探视子女一次的'
                                                   '权利。男方应为女方的探视提供必要的协助。')

        else:
            paragraph = document.add_paragraph('婚后于___（日期）出生的子女由男方抚养。女方每月支付抚养费___元；支付时间是'
                                                   '每月___日之前支付下个月的抚养费；支付方式为银行转账，女方接受抚养费的账号'
                                                   '为:___；直至付到___周岁止，___周岁之后的有关费用双方日后重新协商。')
            paragraph = document.add_paragraph('女方每___（周/月等）有在保证子女身心健康的基础上探视子女一次的权利。男方应'
                                                   '为女方的探视提供必要的协助。')

            paragraph = document.add_paragraph('女方可在____（时间）早上八时接儿子到其居住地，于___（时间）送回王某'
                                                   '居住地，男方应为女方的探视提供必要的协助。如临时或春节探望，可提前'
                                                   '一天与王某协商，达成一致意见后方可探望。')
    for i in range(form['childnum']):
        if form['fangdai']==0:
            paragraph = document.add_paragraph('位于___市___区/县路___号___单元___号，共___平方，登记户主为___的房产，现尚剩余贷'
                                               '款本金___万元，现协商房屋归___方所有，剩余贷款由___方承担。协议经登记生'
                                               '效后，___方应在30日内，配合___方办理贷款主贷人变更手续，以及产权变更手续。')
            paragraph = document.add_paragraph('因办理产权变更登记手续所应支付的一切税费由双方平均承担。')
        else:
            paragraph = document.add_paragraph('位于___市___区/县路___号___单元___号，共___平方，登记户主为___的房产，'
                                               '现协商房屋归___方所有。协议经登记生效后'
                                               '，___方应在30日内，配合___方办理产权变更手续。')
            paragraph = document.add_paragraph('因办理产权变更登记手续所应支付的一切税费由双方平均承担。')
    paragraph = document.add_paragraph('第三条 夫妻共同财产的分割')

    if form['newsaving'] == 0:
        paragraph = document.add_paragraph('婚后各自的户名下的存款，男方名下存款_________元，女方名下存款__________元，'
                                               '经协商婚后女方分的____元，男方分的___元。'
                                               '支付方式：___方向___方在________（日期）之前支付____元。')
    else:
        paragraph = document.add_paragraph('婚后各自的户名下的存款，归各自所有。双方互不补偿。')
    paragraph = document.add_paragraph('第四条 甲乙双方婚前个人财产的确认')

    if form['othersaving']==0:
        paragraph = document.add_paragraph('其他共同财产的分割约定见副本1')
    else:
        paragraph = document.add_paragraph('婚前双方各自的财产归各自所有，男女双方各自的私人生活用品及首饰归各自所有（附清单）')

    paragraph = document.add_paragraph('第五条 债权债务的处理')

    if form['gongtongzhaiwu']==0:
        paragraph = document.add_paragraph('甲乙双方确认在婚姻关系存续期间发生的共同债务由甲乙双方共同偿'
                                               '还：_______年____月____日甲乙双方共同向____________所借债务_____万'
                                               '元(大写：__________________)由甲乙双方各承担50%即_____万元'
                                               '(大写：______________________)。甲乙'
                                               '双方任何一方如对外负有债务的，由负债方自行承担。')
    else:
        paragraph = document.add_paragraph('{双方确认在婚姻关系存续期间没有发生任何'
                                               '共同债权和债务，如任何一方对外负有债权的'
                                               '，无论何时发现，另一方均有权平分；如对外负有债务的，则由负债方自行承担。}')

    paragraph = document.add_paragraph('{离婚后，一方不得干预另一方的生活，不得向第三方泄露另一方的个人隐私及商业秘密，不得有故'
                                       '意损坏另一方名誉的行为，否则承担违约金_________元人民币。}')
    document.save(name)

if __name__ == "__main__":
    form = {
        'sex': 0,
        'name':'张兴宇',
        'childnum': 3,
        'childtoparent':1,
        'tanwangfangshi':1, #0探望 #1短期生活
        'housenum':1,
        'fangdai':0,
        'newsaving':0,
        'othersaving':0,
        'gongtongzhaiwu':0

    }
    main_algo_generate(form,'test34.docx')
