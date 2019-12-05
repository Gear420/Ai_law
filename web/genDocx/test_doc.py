from docx import Document

document = Document()
# paragraph = document.add_paragraph('这是个段落。')
# prior_paragraph = paragraph.insert_paragraph_before('这是前面的段落。')
# document.add_heading('这是个标题')
# #document.add_paragraph('这是一个样式为 ListBullet 的段落', style='ListBullet')

paragraph = document.add_paragraph('男方:{}，男，汉族,{}年{}月{}日生住{},'.format("",""))
document.save("test.docx")