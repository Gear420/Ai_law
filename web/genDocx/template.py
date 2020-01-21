from docx import Document


def main():

    document = Document()
# paragraph = document.add_paragraph('这是个段落。')
# prior_paragraph = paragraph.insert_paragraph_before('这是前面的段落。')
# document.add_heading('这是个标题')
# #document.add_paragraph('这是一个样式为 ListBullet 的段落', style='ListBullet'


    par = document.add_paragraph('这是个段落。')
    par = document.add_paragraph('这是个段落。')
    par = document.add_paragraph('这是个段落。')
    document.save("test.docx")

if __name__ == "__main__":
    main()